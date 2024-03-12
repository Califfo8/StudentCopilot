import whisper
import docx
import json
from pathlib2 import Path
from reader.simpleUI import simpleUI
import ffmpeg
import time
import os
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

class StudentCopilot:
    def __init__(self, config_path):

        f = open(config_path + 'config.json')
        self.data = json.load(f)
        f.close()

        f = open('./appdata/rates.json')
        self.rates = json.load(f)
        f.close()

        self.config_path = config_path + 'config.json'
        self.alfa = 0.125
        self.audio_path = self.data['audio_path'] + "/"
        self.transcription_path = self.data['dest_path']+ "/"
        self.lingua = self.data['lingua']
        self.model_num = self.data['modello']
        self.prompt = self.data['ask_prompt']

        self.start_transcription = -1

    def get_type_of_model(self, model):
        if model == "large":
            return whisper.load_model(model)

        if self.lingua == 2:
            return whisper.load_model(model + ".en")
        elif self.lingua == 1:
            return whisper.load_model(model)
        else:
            print("ERRORE: La lingua impostata non è valida")
            return None

    def get_model_name(self):
        if self.model_num == 1:
            return "tiny"
        elif self.model_num == 2:
            return "base"
        elif self.model_num == 3:
            return "small"
        elif self.model_num == 4:
            return "medium"
        elif self.model_num == 5:
            return "large"
        else:
            print("ERRORE: Il modello selezionato non è valido")
            return None
    def get_model(self):
        return self.get_type_of_model(self.get_model_name())

    # Ritorna la durata dell'audio in secondi
    def get_audio_duration(self, audio_name):
        file_path = self.audio_path + audio_name
        probe = ffmpeg.probe(file_path)
        stream = next((stream for stream in probe['streams'] if stream['codec_type'] == 'audio'), None)
        duration = float(stream['duration'])
        return duration
    def stimate_transcription_time(self, duration):
        rate = self.rates[self.get_model_name()]
        if rate == -1:
            return -1, -1, -1
        s_time_sec = int(duration//rate)
        # Converto in ore:minuti:secondi
        s_time_min = s_time_sec // 60
        s_time_sec %= 60
        s_time_ore = s_time_min // 60
        s_time_min %= 60
        return s_time_ore, s_time_min, s_time_sec

    def compute_new_rate(self, audio_duration):
        if self.start_transcription == -1:
            self.start_transcription = time.time()
            return
        trans_duration = self.start_transcription - time.time()
        if trans_duration == 0:
            print("WARNING: Possibile errore nella stima della velocità di calcolo, l'operazione di stima è stata abortita")
            return
        model_name = self.get_model_name()
        new_rate = audio_duration//trans_duration
        old_rate = self.rates[model_name]
        mean_rate = int(self.alfa * new_rate + (1-self.alfa) * old_rate)

        #Aggiorno il rate
        self.rates[model_name] = mean_rate
        with open("./appdata/rates.json", "w") as outfile:
            json.dump(self.rates, outfile)
        self.start_transcription = -1

    def create_word_file(self, file_name, text, save=True):
        print("Esportando in microsoft word...")
        title = file_name[0:file_name.find('.')]
        # Create a document
        doc = docx.Document()
        # Add a paragraph to the document
        p = doc.add_paragraph()
        # Add a run to the paragraph
        run = p.add_run(title)
        p.style = doc.styles['Heading 2']
        # Add another paragraph
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        # Add a run and format it
        run = p.add_run(text)
        # Save the document
        doc.save(self.transcription_path + title + ".docx")
        print("Trascritto file " + file_name + " con successo!")

    def get_all_files(self):
        elements = os.listdir(self.audio_path)
        only_files = []
        for f in elements:
            path = Path(self.audio_path + f)
            if path.is_file():
                only_files.append(f)
        return only_files

    def print_all_files_names(self, only_files):
        print("I seguenti file verranno trascritti:")
        for f in only_files:
            print("\t - " + f)
    def ask_for_confermation(self, file_array):
        self.print_all_files_names(file_array)
        audio_duration = 0
        for f in file_array:
            audio_duration += self.get_audio_duration(f)
        s_time_ore, s_time_min, s_time_sec = self.stimate_transcription_time(audio_duration)
        if s_time_ore == -1:
            print("Primo utilizzo del modello, la stima verrà calcolata dal secondo utilizzo.")
        else:
            print(
                f"Tempo necessario alla trascrizione stimato [modello: {self.get_model_name()}]: {s_time_ore} ore {s_time_min} minuti {s_time_sec} secondi")

        risp = input("Vuoi continuare? [y/n]")
        while risp != 'y' and risp != 'n':
            risp = input("Vuoi continuare? [y/n]")
        if risp == 'n':
            return -1
        return audio_duration

    def convert_speech_to_text_openai(self, file_name, prompt=""):
            FILE_PATH = self.audio_path + file_name
            audio_duration = self.ask_for_confermation([file_name])
            if audio_duration == -1:
                return

            self.compute_new_rate(audio_duration)
            print("Loading model...")
            model = self.get_model()
            print("Starting transcription...")
            result = model.transcribe(FILE_PATH, verbose=True, fp16=False, initial_prompt=prompt)

            self.create_word_file(file_name, result['text'])
            self.compute_new_rate(audio_duration)

    def convert_all_speech_to_text(self, prompt=""):
        only_files = self.get_all_files()
        audio_duration = self.ask_for_confermation(only_files)
        if audio_duration == -1:
            return

        self.compute_new_rate(audio_duration)
        print("Loading model...")
        model = self.get_model()
        print("Starting transcription...")
        doc = docx.Document()
        for f in only_files:
            file_path = self.audio_path + f
            result = model.transcribe(file_path, verbose=False, fp16=False, initial_prompt=prompt)
            title = f[0:f.find('.')]
            # Add a paragraph to the document
            p = doc.add_paragraph()
            # Add a run to the paragraph
            run = p.add_run(title)
            p.style = doc.styles['Heading 2']
            # Add another paragraph
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            # Add a run and format it
            run = p.add_run(result['text'])
            print("Trascritto " + f)
        # Save the document
        doc.save(self.transcription_path + "Trascrizioni.docx")
        self.compute_new_rate(audio_duration)
        print("Operazione conclusa con successo!")

    def convert_all_to_many(self, prompt=""):
        only_files = self.get_all_files()
        audio_duration = self.ask_for_confermation(only_files)
        if audio_duration == -1:
            return

        self.compute_new_rate(audio_duration)
        print("Loading model...")
        model = self.get_model()
        print("Starting transcription...")
        doc = docx.Document()
        for f in only_files:
            file_path = self.audio_path + f
            result = model.transcribe(file_path, verbose=False, fp16=False, initial_prompt=prompt)
            self.create_word_file(f, result['text'])
        self.compute_new_rate(audio_duration)
        print("Operazione conclusa con successo!")
        
    def options(self):
        scelta = -1
        while scelta != 6:
            sUI = simpleUI()
            sUI.clear_console()
            if self.lingua == 1:
                lingua = "Multilingua"
            else:
                lingua = "Inglese"

            if self.prompt:
                ask_prompt = "SI"
            else:
                ask_prompt = "NO"
            scelta = input("-------------------------OPZIONI----------------------------\n"
                  "Scegli l'impostazione da cambiare:\n"
                  "[1] Percorso file audio: " + self.audio_path + "\n"
                  "[2] Percorso file trascrizioni: " + self.transcription_path + "\n"
                  "[3] Lingua preferita: " + lingua + "\n"
                  "[4] Modello scelto: " + self.get_model_name() + "\n"
                  "[5] Chiedi sempre la frase di contesto: " + ask_prompt + "\n"
                  "[6] Ritorna al menu\n")
            scelta = int(scelta)
            if scelta == 1:
                self.audio_path = input("Inserire il nuovo percorso (completo ed assoluto)'\n")
                self.data['audio_path'] = self.audio_path
            elif scelta == 2:
                self.transcription_path = input("Inserire il nuovo percorso (completo ed assoluto)'\n")
                self.data['dest_path'] = self.transcription_path
            elif scelta == 3:
                self.lingua = input("Scegliere la lingua degli audio su cui si effettuano le trascrizioni:\n"
                                 "[1] Multilingua: rilevamento automatico, italiano incluso. La trascrizione è più lenta\n"
                                 "[2] Inglese: specifico per la lingua, più precisa e veloce per quest'ultima\n ")
                self.lingua = int(self.lingua)
                self.data['lingua'] = int(self.lingua)
            elif scelta == 4:
                self.model_num = input("Scegliere il modello AI da utilizzare per le trascrizioni, più un modello è potente più è preciso e più richiede risorse\n"
                                 "La tabella è un estratto dalla pagina di OpenAI di Whisper, l'ultima colonna indica quanto velocemente l'AI 'riproduce' l'audio\n."
                                 "Per la versione large si consideri che non è presente un modello specializzato sulla trascrizione inglese, vi è solo quello Multilingua.\n"
                                 "Ne segue che la scelta precedente non avrà effetti\n"
                                 "      | Size   | Required VRAM | Relative speed\n"
                                 "  [1] | tiny   |     ~1 GB     |      ~32x\n"
                                 "  [2] | base   |     ~1 GB     |      ~16x\n"
                                 "  [3] | small  |     ~2 GB     |      ~6x\n"
                                 "  [4] | medium |     ~5 GB     |      ~2x\n"
                                 "  [5] | large  |     ~10 GB    |      ~1x\n")
                self.model_num = int(self.model_num)
                self.data['modello'] = self.model_num
            elif scelta == 5:
                self.prompt = input(
                    "Il modello di intelligenza artificiale può utilizzare una breve frase (scritta in inglese) che descriva\n"
                    "il contenuto dell'audio per migliorare la sua precisione. Tuttavia tale inserimento è falcoltativo,\n"
                    "vuoi che ad ogni comando di trascrizione ti venga richiesto di inserire questa frase? [y/n]\n")
                if self.prompt == "y":
                    self.prompt = True
                else:
                    self.prompt = False
                self.data['ask_prompt'] = self.prompt

        with open(self.config_path, "w") as outfile:
            json.dump(self.data, outfile)
        print("Configurazione completata con successo!")